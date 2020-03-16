#-*- coding:utf-8 -*-
import random
import docx
from docx.shared import Pt
import time

def convert2Money(num):
    result = ''
    if num >= 10:
        result = '{}元'.format(int(num / 10))
    if num % 10 != 0:
        result += '{}角'.format(num % 10)
    return result

'''
line_count:行数
begin：起始数字
end：截止数字
each_line_problems：每行问题个数
setting_carry：0: 不含有进位(退位)
               1：都是进位(退位)
               2: 混合
mode:  1  加法
        2  减法
        3  加减法
        4  钱的加减法
'''
def unique_random_engine(line_count=17, begin=2, end=9, each_line_problems=3, setting_carry=0, mode = 1):
    f = docx.Document()
    style = f.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(20)
    s = []
    if mode == 1 or mode == 3 or mode == 4:
        s = [(i, '+', j) for i in range(begin, end + 1) for j in range(begin, end + 1) if i + j < 100]
    if mode == 2 or mode == 3 or mode == 4:
        sub_s = [(i, '-', j) for i in range(begin, end + 1) for j in range(begin, i)]
        s = s + sub_s
    random.shuffle(s)
    index = 0
    if mode == 4:
        each_line_problems = 1
    for count in range(line_count):
        line = ''
        k = 0
        while k < each_line_problems:
            if index >= len(s):
                index = 0
                random.shuffle(s)
            data = s[index]
            if setting_carry is 0:
                if data[1] is '+' and (data[0] % 10 + data[2] % 10 > 9):
                    index = index + 1
                    continue
                if data[1] is '-' and (data[0] % 10 < data[2] % 10):
                    index = index + 1
                    continue
            if setting_carry is 1:
                if data[1] is '+' and (data[0] % 10 + data[2] % 10 <= 9):
                    index = index + 1
                    continue
                if data[1] is '-' and (data[0] % 10 >= data[2] % 10):
                    index = index + 1
                    continue
            if mode == 4:
                line = '{}   {}   {}  =               '.format(convert2Money(data[0]), data[1], convert2Money(data[2]))
            else:
                line = '{}{:<2d}  {}  {:<2d}  =               '.format(line, data[0], data[1], data[2])
            #line = line + str(data[0]) + '  ' + str(data[1]) + '  ' + str(data[2]) + '  =           '
            index = index + 1
            k = k + 1
        print(line)
        f.add_paragraph(line)
    f.save('.\\Problems\\Unique_Random_{}_{}_{}_{}_{}.docx'.format(time.time(), begin, end, setting_carry, mode))

'''
mode:  1  加法
    2  减法
    3  加减法
'''
def unique_smart_engine(line_count=17, begin=2, end=99, each_line_problems=2, mode=1):
    f = docx.Document()
    style = f.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(20)
    s = []

    if mode == 1 or mode == 3:
        s = [(i, '+', j) for i in range(begin, end + 1) for j in range(begin, end + 1) if i + j < end and (i + j) % 10 == 0]
    if mode == 2 or mode == 3:
        sub_s = [(i, '-', j) for i in range(begin, end + 1) for j in range(begin, i) if (i % 10) == (j % 10)]
        s = s + sub_s
    random.shuffle(s)
    index = 0
    for count in range(line_count):
        line = ''
        k = 0
        while k < each_line_problems:
            if index >= len(s):
                index = 0
                random.shuffle(s)
            data = s[index]
            op = random.choice(['+', '-'])
            if data[1] == '+':
                val = data[0] + data[2]
            else:
                val = data[0] - data[2]
            if op == '+':
                d = random.choice(range(1, end + 1 - val))
            else:
                d = random.choice(range(1, val))
            if (op == '-' and d == val) or (d == data[2]):
                continue
            pos = random.choice([0, 1])
            if pos is 0:
                line = '{}{:<2d}  {}  {:<2d}  {}  {:<2d}  =                '.format(line, data[0], op, d, data[1], data[2])
            else:
                line = '{}{:<2d}  {}  {:<2d}  {}  {:<2d}  =                '.format(line, data[0], data[1], data[2], op, d)
            index = index + 1
            k = k + 1
        print(line)
        f.add_paragraph(line)
    f.save('.\\Problems\\Unique_Smart_{}_{}_{}_{}.docx'.format(time.time(), begin, end, mode))



if __name__ == '__main__':
    unique_random_engine(170, begin=10, end=99, setting_carry=1, mode=3)
    #unique_smart_engine(170, 2, 50, 2, 3)
    print('end')
